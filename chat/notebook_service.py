"""
Notebook LLM Service - Mimics Google Notebook LLM functionality
Handles file uploads, notebook creation, and source management
"""
from django.conf import settings
import google.generativeai as genai
import os
import json
from typing import List, Dict, Optional
import tempfile
import base64


class NotebookLLMService:
    """
    Service to manage Notebooks and Sources similar to Google Notebook LLM
    Uses Gemini API for file processing
    """
    
    def __init__(self):
        if settings.GEMINI_API_KEY:
            genai.configure(api_key=settings.GEMINI_API_KEY)
    
    def upload_file_to_gemini(self, file_path: str, mime_type: str = None) -> Optional[genai.types.File]:
        """
        Upload a file to Gemini for processing (like Notebook LLM sources)
        Returns a Gemini File object that can be referenced in chat
        """
        try:
            file = genai.upload_file(path=file_path, mime_type=mime_type)
            print(f"[Notebook LLM] Uploaded file '{file.display_name}' as: {file.uri}")
            return file
        except Exception as e:
            print(f"[Notebook LLM] Upload error: {e}")
            return None
    
    def process_source_content(self, file_obj, content: str = None) -> Dict:
        """
        Process a source file and extract key information
        Similar to how Notebook LLM processes uploaded documents
        """
        source_data = {
            'id': None,
            'name': getattr(file_obj, 'name', 'Unknown'),
            'type': getattr(file_obj, 'type', 'text/plain'),
            'size': getattr(file_obj, 'size', 0),
            'content_preview': (content or '')[:500] if content else '',  # First 500 chars
            'gemini_file_uri': None,
            'extracted_text': content or ''
        }
        
        # If we have actual content, process it with Gemini to create a summary
        if content and len(content) > 100:
            try:
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = f"""Analyze this document and provide:
1. A brief summary (2-3 sentences)
2. Key topics/themes
3. Important facts or data points

Document content (first 5000 chars):
{content[:5000]}
"""
                response = model.generate_content(prompt)
                source_data['ai_summary'] = response.text
            except Exception as e:
                print(f"[Notebook LLM] AI summary error: {e}")
                source_data['ai_summary'] = None
        
        return source_data
    
    def create_notebook_context(self, sources: List[Dict], query: str = "") -> str:
        """
        Create a context string from notebook sources - formatted for AI consumption
        Ensures proper structure that doesn't interfere with <thinking> tags protocol
        Includes FULL content so model can output everything when requested
        """
        if not sources:
            return ""

        context_parts = ["\n=== NOTEBOOK LLM SOURCES ==="]
        context_parts.append("INSTRUCTION: When user asks to read/show/output file content, you MUST output the COMPLETE and FULL content. Do NOT summarize, shorten, or skip any part.\n")

        for source in sources:
            # Use title for URLs, name for files
            source_name = source.get('name') or source.get('title', 'Unknown')
            source_type = source.get('type', 'unknown')
            context_parts.append(f"\n[File: {source_name}]")
            context_parts.append(f"Type: {source_type}")

            # Add URL for URL-type sources
            if source_type == 'url' and source.get('url'):
                context_parts.append(f"URL: {source['url']}")

            context_parts.append(f"Size: {source.get('size', 0)} bytes")

            # Include AI summary if available (short)
            if source.get('ai_summary'):
                context_parts.append(f"Summary: {source['ai_summary'][:200]}")

            # Include FULL content - no truncation
            if source.get('extracted_text'):
                full_content = source['extracted_text']
                context_parts.append(f"\nFULL CONTENT START:")
                context_parts.append(full_content)
                context_parts.append("FULL CONTENT END")

            context_parts.append("")  # Empty line between sources

        context_parts.append("\n=== END NOTEBOOK LLM SOURCES ===\n")
        context_parts.append("QUY TẮC: Suy nghĩ trong <thinking> rồi trả lời. Nếu user yêu cầu đọc file, hãy output TOÀN BỘ nội dung từ sources.")
        context_parts.append(f"Câu hỏi: {query}\n")
        
        return "\n".join(context_parts)
    
    def query_with_notebooks(self, user_message: str, notebook_sources: List[Dict], 
                            conversation_history: List[Dict] = None) -> str:
        """
        Process a query with notebook sources attached
        Returns a grounded response based on the sources
        """
        # Create context from notebooks
        notebook_context = self.create_notebook_context(notebook_sources, user_message)
        
        # Build the full prompt
        full_prompt = f"""{notebook_context}

User Question: {user_message}

Instructions:
1. Answer based ONLY on the provided notebook sources above
2. Cite specific sources using [Source X] format
3. If the answer isn't in the sources, say "I don't find this information in your notebook sources"
4. Be concise but thorough
"""
        
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # Build history if provided
            if conversation_history:
                chat = model.start_chat(history=conversation_history)
                response = chat.send_message(full_prompt)
            else:
                response = model.generate_content(full_prompt)
            
            return response.text
            
        except Exception as e:
            print(f"[Notebook LLM] Query error: {e}")
            return f"Error processing notebook query: {str(e)}"


# Global service instance
notebook_service = NotebookLLMService()


def extract_text_from_file(file_obj) -> str:
    """
    Extract text content from uploaded file
    Supports: txt, md, pdf, docx, doc
    """
    import io
    import re
    
    name = getattr(file_obj, 'name', '').lower()
    content = ''
    
    try:
        if name.endswith(('.txt', '.md', '.csv', '.json')):
            # Text files - read directly
            content = file_obj.read().decode('utf-8')
            
        elif name.endswith('.pdf'):
            # PDF - try pdfplumber first, then PyPDF2
            file_obj.seek(0)
            pdf_bytes = file_obj.read()
            
            # PDF extraction with multiple fallback methods
            content = ""
            errors = []
            
            # Method 1: Try pdfplumber (best for Vietnamese and complex PDFs)
            try:
                import pdfplumber
                pdf_file = io.BytesIO(pdf_bytes)
                text_parts = []
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(f"[Trang {page_num}]\n{page_text}")
                        except Exception as page_err:
                            errors.append(f"Page {page_num}: {str(page_err)[:50]}")
                            continue
                
                if text_parts:
                    content = "\n\n".join(text_parts)
                    content = re.sub(r'\n{3,}', '\n\n', content)
                    content = content.strip()
                    print(f"[PDF] pdfplumber extracted {len(content)} chars from {len(text_parts)} pages")
                    
            except Exception as e1:
                errors.append(f"pdfplumber: {str(e1)[:80]}")
                print(f"[PDF] pdfplumber failed: {e1}")
            
            # Method 2: If pdfplumber failed or got little text, try PyPDF2
            if len(content) < 100:
                try:
                    from PyPDF2 import PdfReader
                    pdf_file = io.BytesIO(pdf_bytes)
                    reader = PdfReader(pdf_file)
                    text_parts = []
                    
                    for page_num, page in enumerate(reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(f"[Trang {page_num}]\n{page_text}")
                        except Exception as page_err:
                            continue
                    
                    if text_parts:
                        content = "\n\n".join(text_parts)
                        content = re.sub(r'\n{3,}', '\n\n', content)
                        content = content.strip()
                        print(f"[PDF] PyPDF2 extracted {len(content)} chars from {len(text_parts)} pages")
                        
                except Exception as e2:
                    errors.append(f"PyPDF2: {str(e2)[:80]}")
                    print(f"[PDF] PyPDF2 failed: {e2}")
            
            # If still no content, mark as unscannable
            if not content or len(content) < 20:
                content = f"[PDF: File scan/ảnh hoặc không đọc được text - {name}]"
                print(f"[PDF] Failed to extract: {'; '.join(errors[:2])}")
            
        elif name.endswith('.docx'):
            # DOCX - use python-docx with enhanced extraction
            try:
                from docx import Document
                from docx.oxml import parse_xml
                
                file_obj.seek(0)
                doc_bytes = file_obj.read()
                doc_file = io.BytesIO(doc_bytes)
                doc = Document(doc_file)
                
                text_parts = []
                
                # 1. Extract from main body paragraphs (all paragraphs, not just text ones)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        # Include style info for structure
                        style = para.style.name if para.style else 'Normal'
                        if style.startswith('Heading'):
                            text_parts.append(f"## {text}")
                        else:
                            text_parts.append(text)
                
                # 2. Extract from tables (including nested tables)
                def extract_table_content(table, depth=0):
                    table_text = []
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            # Get all paragraphs in cell
                            cell_text = []
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    cell_text.append(para.text.strip())
                            # Also check for nested tables
                            for nested_table in cell.tables:
                                nested_text = extract_table_content(nested_table, depth + 1)
                                if nested_text:
                                    cell_text.append(nested_text)
                            if cell_text:
                                row_texts.append(" ".join(cell_text))
                        if row_texts:
                            table_text.append(" | ".join(row_texts))
                    return "\n".join(table_text)
                
                for table in doc.tables:
                    table_content = extract_table_content(table)
                    if table_content:
                        text_parts.append("\n[BẢNG]\n" + table_content + "\n[/BẢNG]\n")
                
                # 3. Extract from headers
                for section in doc.sections:
                    header = section.header
                    if header:
                        header_text = []
                        for para in header.paragraphs:
                            if para.text.strip():
                                header_text.append(para.text.strip())
                        if header_text:
                            text_parts.insert(0, "[HEADER]\n" + "\n".join(header_text) + "\n[/HEADER]\n")
                    
                    # 4. Extract from footers
                    footer = section.footer
                    if footer:
                        footer_text = []
                        for para in footer.paragraphs:
                            if para.text.strip():
                                footer_text.append(para.text.strip())
                        if footer_text:
                            text_parts.append("\n[FOOTER]\n" + "\n".join(footer_text) + "\n[/FOOTER]\n")
                
                # 5. Try to get text from XML for text boxes and other shapes
                try:
                    xml_content = doc.element.xml
                    # Look for text in drawingML shapes (text boxes)
                    import re
                    text_box_pattern = r'<w:t[^>]*>([^<]+)</w:t>'
                    all_texts = re.findall(text_box_pattern, xml_content)
                    # Filter out duplicates already captured
                    existing_text = " ".join(text_parts)
                    for t in all_texts:
                        if t.strip() and len(t.strip()) > 3 and t.strip() not in existing_text:
                            # This might be text box content not in main paragraphs
                            pass  # Skip for now to avoid duplicates
                except:
                    pass
                
                # Combine all content
                content = "\n\n".join(text_parts)
                content = re.sub(r'\n{4,}', '\n\n\n', content)
                content = content.strip() if content else f"[DOCX: File rỗng - {name}]"
                
                # Add metadata
                content = f"[TÀI LIỆU: {name}]\nĐộ dài: {len(content)} ký tự\n\n{content}"
                
            except Exception as e:
                content = f"[DOCX Error: {str(e)} - {name}]"
                
        elif name.endswith('.doc'):
            # Old DOC format - limited support
            try:
                file_obj.seek(0)
                text = file_obj.read().decode('utf-8', errors='ignore')
                text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
                content = text.strip() if text.strip() else f"[DOC: Không thể đọc file cũ {name}. Vui lòng chuyển sang .docx]"
            except Exception as e:
                content = f"[DOC Error: {str(e)} - {name}]"
            
        else:
            content = f"[File: {name}]\n[Binary or unsupported format]"
            
    except Exception as e:
        content = f"[Error reading file {name}: {str(e)}]"
    
    return content


def extract_pdf_content_and_images(file_obj) -> Dict:
    """
    Extract both text and images from PDF using PyMuPDF (fitz)
    Returns dict with 'text' and 'images' keys
    """
    import io
    import base64
    import tempfile
    import os

    name = getattr(file_obj, 'name', '').lower()
    result = {'text': '', 'images': []}

    if not name.endswith('.pdf'):
        return result

    try:
        file_obj.seek(0)
        pdf_bytes = file_obj.read()

        # Try PyMuPDF (fitz) for comprehensive extraction
        try:
            import fitz  # PyMuPDF

            pdf_file = io.BytesIO(pdf_bytes)
            doc = fitz.open(stream=pdf_file, filetype="pdf")

            text_parts = []
            images = []

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Trang {page_num + 1}]\n{page_text}")

                # Extract images from this page
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Convert to base64 for storage
                    image_b64 = base64.b64encode(image_bytes).decode('utf-8')

                    images.append({
                        'page': page_num + 1,
                        'index': img_index,
                        'data': f"data:image/{image_ext};base64,{image_b64}",
                        'size': len(image_bytes)
                    })

            doc.close()

            result['text'] = "\n\n".join(text_parts)
            result['images'] = images

            print(f"[PDF MuPDF] Extracted {len(text_parts)} pages, {len(images)} images")
            return result

        except ImportError:
            print("[PDF] PyMuPDF not available, falling back to text-only extraction")
        except Exception as e:
            print(f"[PDF MuPDF] Error: {e}")

        # Fallback: pdfplumber for text only
        try:
            import pdfplumber
            pdf_file = io.BytesIO(pdf_bytes)
            text_parts = []

            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"[Trang {page_num}]\n{page_text}")

            result['text'] = "\n\n".join(text_parts)
            print(f"[PDF] pdfplumber extracted {len(text_parts)} pages")

        except Exception as e:
            print(f"[PDF] pdfplumber failed: {e}")
            result['text'] = "[PDF: Không thể trích xuất text]"

    except Exception as e:
        print(f"[PDF] Error: {e}")
        result['text'] = f"[PDF Error: {str(e)}]"

    return result


def process_uploaded_files_for_notebook(files) -> List[Dict]:
    """
    Process multiple uploaded files and return list of source dictionaries
    For PDFs: extracts both text and images
    """
    sources = []

    for file in files:
        print(f"[Notebook LLM] Processing file: {file.name}, size: {file.size}")

        # Check if PDF for enhanced extraction
        is_pdf = file.name.lower().endswith('.pdf')

        if is_pdf:
            # Use enhanced PDF extraction with images
            pdf_result = extract_pdf_content_and_images(file)
            content = pdf_result['text']
            pdf_images = pdf_result['images']
            print(f"[Notebook LLM] PDF extracted: {len(content)} chars, {len(pdf_images)} images")
        else:
            # Use regular text extraction
            content = extract_text_from_file(file)
            pdf_images = []

        print(f"[Notebook LLM] Extracted content length: {len(content)}, preview: {content[:100] if content else 'EMPTY'}")

        # Process through Notebook LLM service
        source = notebook_service.process_source_content(file, content)

        # Add images to source if PDF
        if pdf_images:
            source['pdf_images'] = pdf_images
            source['has_images'] = True

        print(f"[Notebook LLM] Source created with extracted_text length: {len(source.get('extracted_text', ''))}")

        # Also try to upload to Gemini for advanced processing (only supported file types)
        # Note: Gemini doesn't support DOCX MIME type (application/vnd.openxmlformats-officedocument.wordprocessingml.document)
        supported_gemini_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.txt', '.md', '.csv', '.json']
        file_ext = os.path.splitext(file.name)[1].lower()
        
        if file_ext in supported_gemini_extensions:
            try:
                # Save to temp file for Gemini upload
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                    file.seek(0)
                    tmp.write(file.read())
                    tmp_path = tmp.name

                # Upload to Gemini with explicit MIME type for PDF
                mime_type = None
                if file_ext == '.pdf':
                    mime_type = 'application/pdf'
                elif file_ext in ['.jpg', '.jpeg']:
                    mime_type = 'image/jpeg'
                elif file_ext == '.png':
                    mime_type = 'image/png'
                elif file_ext == '.gif':
                    mime_type = 'image/gif'
                elif file_ext == '.webp':
                    mime_type = 'image/webp'
                elif file_ext == '.txt':
                    mime_type = 'text/plain'
                elif file_ext == '.md':
                    mime_type = 'text/plain'
                elif file_ext == '.csv':
                    mime_type = 'text/csv'
                elif file_ext == '.json':
                    mime_type = 'application/json'
                
                gemini_file = notebook_service.upload_file_to_gemini(tmp_path, mime_type=mime_type)
                if gemini_file:
                    source['gemini_file_uri'] = gemini_file.uri
                    source['id'] = gemini_file.name
                    print(f"[Notebook LLM] Gemini upload successful: {gemini_file.uri}")

                # Clean up temp file
                os.unlink(tmp_path)

            except Exception as e:
                print(f"[Notebook LLM] Gemini upload failed for {file.name}: {e}")
        else:
            print(f"[Notebook LLM] Skipping Gemini upload for {file.name} - unsupported file type. Using extracted text only.")

        sources.append(source)

        # Reset file pointer
        file.seek(0)

    return sources
